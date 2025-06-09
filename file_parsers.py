import pandas as pd
from docx import Document
import json
from typing import List, Dict, Optional, Union
import re
import io


def parse_kpi_excel(file_content: bytes) -> List[Dict[str, str]]:
    """
    Parse KPI definitions from Excel file.
    Expected columns: KPI Name, Description, Formula (optional), Target (optional), Category (optional)
    """
    try:
        df = pd.read_excel(io.BytesIO(file_content))
        
        # Normalize column names (case-insensitive)
        df.columns = df.columns.str.strip().str.lower()
        
        # Map possible column variations
        column_mapping = {
            'kpi name': ['kpi', 'name', 'metric', 'indicator', 'kpi_name'],
            'description': ['desc', 'definition', 'description', 'details'],
            'formula': ['calculation', 'formula', 'logic', 'calc'],
            'target': ['target', 'goal', 'threshold', 'benchmark'],
            'category': ['category', 'group', 'department', 'area', 'type']
        }
        
        # Find matching columns
        found_columns = {}
        for standard_col, variations in column_mapping.items():
            for col in df.columns:
                if any(var in col for var in variations):
                    found_columns[standard_col] = col
                    break
        
        # Must have at least name and description
        if 'kpi name' not in found_columns or 'description' not in found_columns:
            # Try to use first two columns as name and description
            if len(df.columns) >= 2:
                found_columns['kpi name'] = df.columns[0]
                found_columns['description'] = df.columns[1]
            else:
                raise ValueError("Excel file must contain at least KPI Name and Description columns")
        
        kpi_list = []
        for _, row in df.iterrows():
            kpi = {
                "name": str(row[found_columns['kpi name']]).strip(),
                "description": str(row[found_columns['description']]).strip()
            }
            
            # Add optional fields if present
            if 'formula' in found_columns and pd.notna(row[found_columns['formula']]):
                kpi["formula"] = str(row[found_columns['formula']]).strip()
            
            if 'target' in found_columns and pd.notna(row[found_columns['target']]):
                kpi["target"] = str(row[found_columns['target']]).strip()
                
            if 'category' in found_columns and pd.notna(row[found_columns['category']]):
                kpi["category"] = str(row[found_columns['category']]).strip()
            
            # Skip empty rows
            if kpi["name"] and kpi["name"].lower() not in ['nan', 'none', '']:
                kpi_list.append(kpi)
        
        return kpi_list
        
    except Exception as e:
        raise ValueError(f"Error parsing KPI Excel file: {str(e)}")


def parse_kpi_word(file_content: bytes) -> List[Dict[str, str]]:
    """
    Parse KPI definitions from Word document.
    Expects bulleted or numbered lists with KPI descriptions.
    """
    try:
        doc = Document(file_content)
        kpi_list = []
        
        # Extract all paragraphs
        text_content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_content.append(text)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_content.append(text)
        
        # Parse KPIs - look for patterns
        current_kpi = None
        for text in text_content:
            # Check if this looks like a KPI name (often starts with bullet, number, or is in caps)
            if re.match(r'^[\d\-•▪►*]+\s*(.+)', text) or text.isupper() or ':' in text:
                # Extract KPI name
                kpi_name = re.sub(r'^[\d\-•▪►*]+\s*', '', text)
                kpi_name = kpi_name.split(':')[0].strip()
                
                # Save previous KPI if exists
                if current_kpi and current_kpi.get("name"):
                    kpi_list.append(current_kpi)
                
                # Start new KPI
                current_kpi = {
                    "name": kpi_name,
                    "description": text.split(':', 1)[1].strip() if ':' in text else ""
                }
            elif current_kpi:
                # This is likely additional description for current KPI
                if current_kpi.get("description"):
                    current_kpi["description"] += " " + text
                else:
                    current_kpi["description"] = text
        
        # Don't forget the last KPI
        if current_kpi and current_kpi.get("name"):
            kpi_list.append(current_kpi)
        
        # If no KPIs found with pattern matching, treat each paragraph as a KPI
        if not kpi_list:
            for i, text in enumerate(text_content):
                if len(text) > 10:  # Skip very short lines
                    kpi_list.append({
                        "name": f"KPI {i+1}",
                        "description": text
                    })
        
        return kpi_list
        
    except Exception as e:
        raise ValueError(f"Error parsing KPI Word file: {str(e)}")


def parse_data_dictionary_excel(file_content: bytes) -> Dict[str, Dict[str, Dict[str, str]]]:
    """
    Parse data dictionary from Excel file.
    Expected columns: Table Name, Column Name, Description, Data Type, Example Values, Business Rules
    """
    try:
        df = pd.read_excel(io.BytesIO(file_content))
        
        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()
        
        # Map possible column variations
        column_mapping = {
            'table_name': ['table', 'table_name', 'table name', 'entity'],
            'column_name': ['column', 'field', 'column_name', 'column name', 'field_name', 'attribute'],
            'description': ['description', 'desc', 'definition', 'business_description'],
            'data_type': ['type', 'data_type', 'datatype', 'data type', 'field_type'],
            'example': ['example', 'sample', 'example_values', 'sample_values'],
            'rules': ['rules', 'constraints', 'business_rules', 'validation', 'notes']
        }
        
        # Find matching columns
        found_columns = {}
        for standard_col, variations in column_mapping.items():
            for col in df.columns:
                if any(var in col for var in variations):
                    found_columns[standard_col] = col
                    break
        
        # Must have at least table, column, and description
        required = ['table_name', 'column_name', 'description']
        for req in required:
            if req not in found_columns:
                # Try to use first three columns
                if len(df.columns) >= 3:
                    found_columns['table_name'] = df.columns[0]
                    found_columns['column_name'] = df.columns[1]
                    found_columns['description'] = df.columns[2]
                else:
                    raise ValueError("Data dictionary must contain at least Table Name, Column Name, and Description")
        
        data_dictionary = {}
        
        for _, row in df.iterrows():
            table_name = str(row[found_columns['table_name']]).strip()
            column_name = str(row[found_columns['column_name']]).strip()
            
            # Skip empty rows
            if not table_name or table_name.lower() in ['nan', 'none', '']:
                continue
                
            if table_name not in data_dictionary:
                data_dictionary[table_name] = {}
            
            column_info = {
                "description": str(row[found_columns['description']]).strip()
            }
            
            # Add optional fields
            if 'data_type' in found_columns and pd.notna(row[found_columns['data_type']]):
                column_info["type"] = str(row[found_columns['data_type']]).strip()
                
            if 'example' in found_columns and pd.notna(row[found_columns['example']]):
                column_info["example"] = str(row[found_columns['example']]).strip()
                
            if 'rules' in found_columns and pd.notna(row[found_columns['rules']]):
                column_info["rules"] = str(row[found_columns['rules']]).strip()
            
            data_dictionary[table_name][column_name] = column_info
        
        return data_dictionary
        
    except Exception as e:
        raise ValueError(f"Error parsing data dictionary Excel file: {str(e)}")


def parse_data_dictionary_csv(file_content: bytes) -> Dict[str, Dict[str, Dict[str, str]]]:
    """
    Parse data dictionary from CSV file.
    Uses same logic as Excel parser.
    """
    try:
        # Try different encodings
        from io import StringIO
        df = None
        for encoding in ['utf-8', 'latin-1', 'iso-8859-1']:
            try:
                # Convert bytes to string and create StringIO object
                text_content = file_content.decode(encoding)
                df = pd.read_csv(StringIO(text_content))
                break
            except UnicodeDecodeError:
                continue
        
        if df is None:
            raise ValueError("Could not decode CSV file with common encodings")
        
        # Apply same parsing logic as Excel parser but with the dataframe we just created
        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()
        
        # Map possible column variations
        column_mapping = {
            'table_name': ['table', 'table_name', 'table name', 'entity'],
            'column_name': ['column', 'field', 'column_name', 'column name', 'field_name', 'attribute'],
            'description': ['description', 'desc', 'definition', 'business_description'],
            'data_type': ['type', 'data_type', 'datatype', 'data type', 'field_type'],
            'example': ['example', 'sample', 'example_values', 'sample_values'],
            'rules': ['rules', 'constraints', 'business_rules', 'validation', 'notes']
        }
        
        # Find matching columns
        found_columns = {}
        for standard_col, variations in column_mapping.items():
            for col in df.columns:
                if any(var in col for var in variations):
                    found_columns[standard_col] = col
                    break
        
        # Must have at least table, column, and description
        required = ['table_name', 'column_name', 'description']
        for req in required:
            if req not in found_columns:
                # Try to use first three columns
                if len(df.columns) >= 3:
                    found_columns['table_name'] = df.columns[0]
                    found_columns['column_name'] = df.columns[1]
                    found_columns['description'] = df.columns[2]
                else:
                    raise ValueError("Data dictionary must contain at least Table Name, Column Name, and Description")
        
        data_dictionary = {}
        
        for _, row in df.iterrows():
            table_name = str(row[found_columns['table_name']]).strip()
            column_name = str(row[found_columns['column_name']]).strip()
            
            # Skip empty rows
            if not table_name or table_name.lower() in ['nan', 'none', '']:
                continue
                
            if table_name not in data_dictionary:
                data_dictionary[table_name] = {}
            
            column_info = {
                "description": str(row[found_columns['description']]).strip()
            }
            
            # Add optional fields
            if 'data_type' in found_columns and pd.notna(row[found_columns['data_type']]):
                column_info["type"] = str(row[found_columns['data_type']]).strip()
                
            if 'example' in found_columns and pd.notna(row[found_columns['example']]):
                column_info["example"] = str(row[found_columns['example']]).strip()
                
            if 'rules' in found_columns and pd.notna(row[found_columns['rules']]):
                column_info["rules"] = str(row[found_columns['rules']]).strip()
            
            data_dictionary[table_name][column_name] = column_info
        
        return data_dictionary
        
    except Exception as e:
        raise ValueError(f"Error parsing data dictionary CSV file: {str(e)}")


def validate_kpi_list(kpi_list: List[Dict[str, str]]) -> bool:
    """Validate that KPI list has required structure"""
    if not kpi_list:
        return False
    
    for kpi in kpi_list:
        if not isinstance(kpi, dict):
            return False
        if 'name' not in kpi or 'description' not in kpi:
            return False
        if not kpi['name'] or not kpi['description']:
            return False
    
    return True


def validate_data_dictionary(data_dict: Dict[str, Dict[str, Dict[str, str]]]) -> bool:
    """Validate that data dictionary has required structure"""
    if not data_dict:
        return False
    
    for table_name, columns in data_dict.items():
        if not isinstance(columns, dict):
            return False
        for col_name, col_info in columns.items():
            if not isinstance(col_info, dict):
                return False
            if 'description' not in col_info:
                return False
    
    return True